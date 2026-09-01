"""
Helper script to update progress at the top of MarketScan_Ultimate_Masterplan_2026-09-01.docx
"""
import docx
from docx.shared import Pt, RGBColor
import datetime
import os
import sys

DOCX_PATH = r"C:\Users\hthur\Downloads\MarketScan_Ultimate_Masterplan_2026-09-01.docx"
MD_PATH = r"C:\Users\hthur\.gemini\antigravity\brain\05d32010-0a00-466d-a39c-7cbb0f60703a\extracted_masterplan.md"

ALL_PHASES = [
    ("PHASE 0.1", "Freeze semantics and create migration branch"),
    ("PHASE 0.2", "Fix P0 production correctness"),
    ("PHASE 0.3", "Security hardening"),
    ("PHASE 1", "Security Master v2"),
    ("PHASE 2", "Provenance + Data Quality platform"),
    ("PHASE 3", "MasterRank v2 challenger"),
    ("PHASE 4", "Event + Analyst revision engines"),
    ("PHASE 5", "SetupState + Risk v1 shadow"),
    ("PHASE 6", "Decision API v2"),
    ("PHASE 7", "UI/UX v2 foundation"),
    ("PHASE 8", "Stock page + cross-product migration"),
    ("PHASE 9", "AI Research v2"),
    ("PHASE 10", "Research engine + calibration"),
    ("PHASE 11", "Portfolio construction v2"),
    ("PHASE 12", "Final cutover")
]

def update_docx_progress(completed_phases: list[dict], current_phase: str, summary_stats: dict):
    with open(MD_PATH, "r", encoding="utf-8") as f:
        lines = [line.strip() for line in f.readlines()]

    doc = docx.Document()

    # Lines 0 to 5 are title and header metadata
    for line in lines[:6]:
        p = doc.add_paragraph(line)
        if line == "MarketScan":
            p.style.font.size = Pt(14)
            p.style.font.bold = True
        elif line == "ULTIMATE MASTERPLAN":
            p.style.font.size = Pt(18)
            p.style.font.bold = True
            p.style.font.color.rgb = RGBColor(0, 50, 150)

    doc.add_paragraph("")

    # Progress Section
    p = doc.add_paragraph("=== EXEKVERINGSSTATUS OCH PROGRESS RAPPORT ===")
    p.style.font.bold = True
    p.style.font.size = Pt(12)
    p.style.font.color.rgb = RGBColor(0, 50, 150)

    now_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    p = doc.add_paragraph(
        f"Status: {summary_stats.get('status', 'Pågående')} | "
        f"Aktuell fas: {current_phase} | "
        f"Klara faser: {len(completed_phases)} / {len(ALL_PHASES)} | "
        f"Senast uppdaterad: {now_str}"
    )
    p.style.font.bold = True
    p.style.font.size = Pt(10)
    p.style.font.color.rgb = RGBColor(0, 50, 150)

    doc.add_paragraph("")
    p = doc.add_paragraph("Genomförda faser och Detaljerad status:")
    p.style.font.bold = True
    p.style.font.size = Pt(10.5)
    p.style.font.color.rgb = RGBColor(20, 100, 30)

    completed_ids = {c["phase"] for c in completed_phases}

    if not completed_phases:
        p = doc.add_paragraph("  (Inga faser helt klarmarkerade ännu - startar Fas 0.1)")
        p.style.font.size = Pt(9)
        p.style.font.color.rgb = RGBColor(120, 120, 120)
    else:
        for comp in completed_phases:
            p = doc.add_paragraph(f"  [X] {comp['phase']}: {comp['title']} [KLAR {comp.get('time', '')}]")
            p.style.font.bold = True
            p.style.font.size = Pt(9.5)
            p.style.font.color.rgb = RGBColor(20, 120, 40)
            for detail in comp.get("details", []):
                p_d = doc.add_paragraph(f"      - {detail}")
                p_d.style.font.size = Pt(9)
            for test in comp.get("tests", []):
                p_t = doc.add_paragraph(f"      + Test/Gate: {test}")
                p_t.style.font.size = Pt(8.5)
                p_t.style.font.color.rgb = RGBColor(50, 100, 50)

    doc.add_paragraph("")
    p = doc.add_paragraph("Kvarstående faser att genomföra:")
    p.style.font.bold = True
    p.style.font.size = Pt(10.5)
    p.style.font.color.rgb = RGBColor(150, 80, 0)

    for pid, title in ALL_PHASES:
        if pid not in completed_ids:
            p = doc.add_paragraph(f"  [ ] {pid}: {title}")
            p.style.font.size = Pt(9.5)
            p.style.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_paragraph("")
    p = doc.add_paragraph("=== SLUT PÅ EXEKVERINGSSTATUS ===")
    p.style.font.size = Pt(9)
    p.style.font.color.rgb = RGBColor(120, 120, 120)
    doc.add_paragraph("")

    # Rest of the original document
    for line in lines[6:]:
        p = doc.add_paragraph(line)
        if line.startswith("PHASE ") or (len(line) > 3 and line[:2].isdigit() and line[2] == "."):
            p.style.font.bold = True

    doc.save(DOCX_PATH)
    print(f"Successfully updated docx progress at {DOCX_PATH} ({now_str}) with {len(doc.paragraphs)} paragraphs.")

if __name__ == "__main__":
    update_docx_progress(
        completed_phases=[],
        current_phase="PHASE 0.1: Freeze semantics and create migration branch",
        summary_stats={"status": "Initierar exekvering"}
    )
